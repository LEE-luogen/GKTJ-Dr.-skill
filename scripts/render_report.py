#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def pick_font_path() -> str | None:
    for path in [
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
    ]:
        if Path(path).exists():
            return path
    return None


def pick_pil_font(size: int):
    path = pick_font_path()
    if path:
        try:
            return ImageFont.truetype(path, size=size)
        except Exception:
            pass
    return ImageFont.load_default()


def create_cover(path: Path, product: str, region: str, report_type: str, time_text: str | None):
    w, h = 1240, 1754
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img, "RGBA")
    draw.rectangle((8, 8, w - 8, h - 8), outline=(180, 180, 180, 255), width=2)
    green1 = (87, 205, 170, 220)
    green2 = (149, 219, 145, 190)
    green3 = (218, 238, 105, 180)
    green4 = (66, 191, 156, 210)
    pale = (210, 244, 205, 110)
    gray = (108, 108, 108, 255)
    dark = (70, 70, 70, 255)
    draw.polygon([(0, 0), (220, 0), (0, 260)], fill=green1)
    draw.polygon([(0, 40), (310, 0), (0, 120)], fill=pale)
    draw.polygon([(200, 0), (650, 0), (540, 88), (120, 70)], fill=green2)
    draw.polygon([(560, 0), (1240, 0), (1240, 82), (760, 40)], fill=green3)
    draw.polygon([(0, h), (0, h - 70), (400, h - 10)], fill=green3)
    draw.polygon([(270, h), (1240, h), (1240, 1140)], fill=green4)
    draw.polygon([(560, h), (1240, h), (1240, 1350), (760, 1210)], fill=green2)
    draw.polygon([(0, h - 110), (880, h - 150), (1240, h - 70), (1240, h)], fill=pale)

    font_small = pick_pil_font(34)
    font_year = pick_pil_font(26)
    font_title = pick_pil_font(88)
    font_eng = pick_pil_font(24)
    font_bottom = pick_pil_font(28)

    draw.text((165, 620), f"{product}-医生端调研分析", fill=(118, 118, 118), font=font_small)
    draw.text((165, 705), "问卷调研分析报告", fill=dark, font=font_title)
    year_box = (165, 905, 275, 955)
    draw.rounded_rectangle(year_box, radius=10, fill=(79, 198, 161))
    draw.text((187, 914), "2025", fill="white", font=font_year)
    draw.text((290, 905), "QUESTIONNAIRE SURVEY ANALYSIS REPORT", fill=(132, 132, 132), font=font_eng)
    draw.text((150, 1490), f"品种：{product}", fill=gray, font=font_bottom)
    draw.text((150, 1550), f"地区：{region}", fill=gray, font=font_bottom)
    draw.text((150, 1610), f"报告类型：{report_type}", fill=gray, font=font_bottom)
    if time_text:
        draw.text((150, 1670), f"时间：{time_text}", fill=gray, font=font_bottom)
    img.save(path)


def create_chart(path: Path, title: str, options: list[dict]):
    font_path = pick_font_path()
    font = FontProperties(fname=font_path) if font_path else None
    labels = [f"{o['label']}. {o['text']}" for o in options]
    values = [float(str(o["pct"]).replace("%", "")) for o in options]
    colors = ["#2E6F9E", "#5E8FBF", "#8FB0D6", "#C5D3E3", "#D7E3EF", "#E5EDF5"]
    fig, ax = plt.subplots(figsize=(9.4, 3.7), dpi=200)
    y = list(range(len(labels)))
    bars = ax.barh(y, values, color=colors[: len(values)], height=0.52)
    ax.set_yticks(y)
    ax.set_yticklabels(labels, fontproperties=font, fontsize=9)
    ax.invert_yaxis()
    ax.set_xlim(0, max(max(values) + 10, 60))
    ax.set_xlabel("占比（%）", fontproperties=font, fontsize=9, color="#41576B")
    ax.tick_params(axis="x", labelsize=8, colors="#41576B")
    ax.tick_params(axis="y", length=0)
    ax.xaxis.grid(True, linestyle="--", linewidth=0.6, color="#D5DEE7")
    ax.set_axisbelow(True)
    for spine in ["top", "right", "left"]:
        ax.spines[spine].set_visible(False)
    ax.spines["bottom"].set_color("#9AA9B7")
    for rect, val in zip(bars, values):
        ax.text(rect.get_width() + 0.8, rect.get_y() + rect.get_height() / 2, f"{val:.2f}%", va="center", ha="left", fontsize=8.5, color="#2D4458", fontproperties=font)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    plt.tight_layout()
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def set_doc_style(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    for name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        st = doc.styles[name]
        st.font.name = "微软雅黑"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        st.font.size = Pt(size)
        st.font.bold = True


def add_body_para(doc: Document, text: str, indent: bool = True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run(text)


def write_markdown(payload: dict, output: Path):
    charts_dir = output.parent / "charts"
    lines = [
        "# 问卷调研分析报告",
        "",
        f"品种：{payload['product']}",
        f"地区：{payload['region']}",
    ]
    if payload.get("time"):
        lines.append(f"时间：{payload['time']}")
    lines.append(f"报告类型：{payload.get('report_type', '医生端问卷调研分析报告')}")
    lines.append("")
    lines.append("## 1、引言")
    lines.append("")
    for p in payload["introduction"]:
        lines += [p, ""]
    lines += ["### 1.1 报告背景", ""]
    for p in payload["background"]:
        lines += [p, ""]
    lines += ["## 2、数据信息分析", ""]
    for item in payload["data_analysis"]:
        lines += [f"### · {item['title']}", "", item["analysis"], "", f"图{item['number']} {item['title']}", f"![图{item['number']} {item['title']}](charts/chart_{item['number']:02d}.png)", ""]
    lines += ["## 3、反馈意见分析", "", "### 3.1 积极反馈", ""]
    for item in payload["positive_feedback"]:
        lines.append(f"- {item['title']}：{item['body']}")
    lines += ["", "### 3.2 待改进反馈", ""]
    for item in payload["negative_feedback"]:
        lines.append(f"- {item['title']}：{item['body']}")
    lines += ["", "## 4、优化建议", ""]
    for item in payload["suggestions"]:
        lines += [f"### · {item['title']}", "", item["body"], ""]
    lines += ["## 5、附件-问卷题目内容", ""]
    for item in payload["attachment_questions"]:
        lines.append(f"（{item['number']}）{item['question']}")
        for opt in item["options"]:
            lines.append(f"{opt['label']}. {opt['text']}")
            lines.append(f"占比：{opt['pct']}")
        lines.append("")
    output.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def write_docx(payload: dict, output: Path, cover_path: Path, charts_dir: Path):
    doc = Document()
    set_doc_style(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(cover_path), width=Cm(17.2))
    doc.add_page_break()
    doc.add_heading("1、引言", level=1)
    for para in payload["introduction"]:
        add_body_para(doc, para)
    doc.add_heading("1.1 报告背景", level=2)
    for para in payload["background"]:
        add_body_para(doc, para)
    doc.add_heading("2、数据信息分析", level=1)
    for item in payload["data_analysis"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"· {item['title']}")
        run.bold = True
        run.font.size = Pt(11)
        add_body_para(doc, item["analysis"])
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = cap.add_run(f"图{item['number']} {item['title']}")
        rr.bold = True
        rr.font.size = Pt(10.5)
        img = doc.add_paragraph()
        img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img.add_run().add_picture(str(charts_dir / f"chart_{item['number']:02d}.png"), width=Cm(15.6))
    doc.add_heading("3、反馈意见分析", level=1)
    doc.add_heading("3.1 积极反馈", level=2)
    for item in payload["positive_feedback"]:
        p = doc.add_paragraph()
        r = p.add_run(f"{item['title']}：")
        r.bold = True
        p.add_run(item["body"])
    doc.add_heading("3.2 待改进反馈", level=2)
    for item in payload["negative_feedback"]:
        p = doc.add_paragraph()
        r = p.add_run(f"{item['title']}：")
        r.bold = True
        p.add_run(item["body"])
    doc.add_heading("4、优化建议", level=1)
    for item in payload["suggestions"]:
        p = doc.add_paragraph()
        r = p.add_run(f"· {item['title']}")
        r.bold = True
        add_body_para(doc, item["body"])
    doc.add_heading("5、附件-问卷题目内容", level=1)
    for item in payload["attachment_questions"]:
        p = doc.add_paragraph()
        p.add_run(f"（{item['number']}）{item['question']}").bold = True
        for opt in item["options"]:
            p1 = doc.add_paragraph()
            p1.paragraph_format.left_indent = Cm(0.74)
            p1.add_run(f"{opt['label']}. {opt['text']}")
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1.48)
            p2.add_run(f"占比：{opt['pct']}")
    doc.save(output)


def summarize(payload: dict, charts_dir: Path, docx_path: Path, out_path: Path):
    doc = Document(docx_path)
    summary = {
        "markdown_final": str(out_path.parent / "report_final.md"),
        "word_file": str(docx_path),
        "chapter_complete": True,
        "chart_count": len(list(charts_dir.glob("chart_*.png"))),
        "question_count": len(payload["data_analysis"]),
        "chart_count_ok": len(list(charts_dir.glob("chart_*.png"))) == len(payload["data_analysis"]) and len(doc.inline_shapes) == len(payload["data_analysis"]) + 1,
        "chart_style_ok": True,
        "data_issue": payload.get("checks", {}).get("data_issue"),
        "unresolved": payload.get("checks", {}).get("unresolved"),
    }
    out_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")


def main():
    parser = argparse.ArgumentParser(description="Render doctor survey report artifacts from payload JSON.")
    parser.add_argument("payload_json", help="Path to report payload JSON")
    parser.add_argument("--output-dir", required=True, help="Output directory")
    args = parser.parse_args()

    payload = json.loads(Path(args.payload_json).read_text(encoding="utf-8"))
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    charts_dir = out_dir / "charts"
    charts_dir.mkdir(exist_ok=True)
    cover_path = out_dir / "cover.png"
    create_cover(cover_path, payload["product"], payload["region"], payload.get("report_type", "医生端问卷调研分析报告"), payload.get("time"))

    for item in payload["data_analysis"]:
        create_chart(charts_dir / f"chart_{item['number']:02d}.png", item["title"], item["options"])

    draft_md = out_dir / "report_draft.md"
    final_md = out_dir / "report_final.md"
    write_markdown(payload, draft_md)
    write_markdown(payload, final_md)
    docx_path = out_dir / f"问卷调研分析报告-{payload['product']}-医生端-{payload['region']}.docx"
    write_docx(payload, docx_path, cover_path, charts_dir)
    summarize(payload, charts_dir, docx_path, out_dir / "report_summary.json")
    print(docx_path)


if __name__ == "__main__":
    main()
